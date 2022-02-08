import pandas as pd
import numpy as np
import docx
from datetime import date

# Adds periods in middle of data to span entire page
# FIXME: Character width is variable, lines are of different length
def addDots(str,index):
    MAX_CHAR = 125 # max chars on one line in Word at 12pt font
    dots_to_add = MAX_CHAR-len(str)
    dots_str = ""

    for i in range(dots_to_add):
        dots_str += "."

    new_str = str[0:index+1] + dots_str + str[index+1:75] # inserts dots in middle
    return new_str

source = 'test_data.xlsx'
services = ["Merit Drench", "Insect Control", "Dormant Oil"]
book = pd.read_excel(source)
report = docx.Document()
year = date.today().year

# book.shape returns a tuple with dimensions of book
rows = book.shape[0]

for i in range(1,rows):
    curr_row = book.iloc[i]
    curr_string = ""

    property_name = curr_row[0]
    curr_string += (property_name + ": ")

    num_trees = curr_row[1]
    if num_trees == 0:
        curr_string += "none"
        curr_string = addDots(curr_string, len(property_name))
        report.add_paragraph(curr_string)
        continue
    elif np.isnan(num_trees):
        curr_string += "N/A"
        report.add_paragraph(curr_string)
        continue
    else:
        curr_string += ("("+str(num_trees)+")")

    flag = curr_row[2]
    curr_string += ("("+str(flag)+")")

    price = curr_row[3]
    if price == "annualized":
        curr_string += "annualized"
        CHAR_BACKWARD = 11; # length of "annualized" + 1
        curr_string = addDots(curr_string, len(curr_string) - CHAR_BACKWARD)
        report.add_paragraph(curr_string)
        continue
    else:
        SALES_TAX_RATE = .0825
        pre_tax = round(price/(1+SALES_TAX_RATE),2)
        tax = round(price-pre_tax,2)

        # "%0.2f" % (num) formats num to exactly two decimal places
        money_str = "$" + "%0.2f" % (pre_tax) + " + $" + "%0.2f" % (tax) + " = $" + "%0.2f" % (price);
        curr_string += money_str
        index = curr_string.find("$")-1
        curr_string = addDots(curr_string, index)
        report.add_paragraph(curr_string)

report_name = str(services[0]).lower() +  " " + str(year % 100 + 1) + ".docx"
report.save(report_name)
print("Report " + report_name + " successfully created")
